from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE = Path(r"C:\Users\Edv\Documents\New project")
PROJECT = WORKSPACE / "ANONYMUS_LINK"
SOURCE = PROJECT / "ANONYMUS_diagrams.md"
OUTPUT = WORKSPACE / "ANONYMUS_diagrams_copyable.docx"


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after, color in [
        ("Heading 1", 20, 20, 6, "000000"),
        ("Heading 2", 16, 18, 6, "000000"),
        ("Heading 3", 14, 16, 4, "434343"),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    code = doc.styles.add_style("Mermaid Code", 1)
    code.font.name = "Courier New"
    code.font.size = Pt(8)
    code.font.color.rgb = RGBColor(25, 25, 25)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.left_indent = Inches(0.12)
    code.paragraph_format.right_indent = Inches(0.12)

    caption = doc.styles.add_style("Small Note", 1)
    caption.font.name = "Arial"
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor(85, 85, 85)
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.line_spacing = 1.15


def parse_mermaid_sections(markdown):
    pattern = re.compile(
        r"##\s+\d+\.\s+(.+?)\n\n```mermaid\n(.*?)\n```",
        re.DOTALL,
    )
    return [(title.strip(), code.strip("\n")) for title, code in pattern.findall(markdown)]


def add_title_block(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("ANONYMUS Project Diagrams")
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph(style="Small Note")
    subtitle.add_run(
        "Copyable Mermaid sources for the system architecture, class, and use case diagrams."
    )

    table = doc.add_table(rows=2, cols=2)
    table.allow_autofit = False
    widths = [Inches(1.55), Inches(4.95)]
    labels = [
        ("Project", "ANONYMUS / Women Safety App"),
        ("Source", "Generated from ANONYMUS_LINK/lib and ANONYMUS_diagrams.md"),
    ]
    for row, (label, value) in zip(table.rows, labels):
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], value)

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F8F9FA")
            tc_pr.append(shd)


def add_copyable_code_block(doc, code):
    for line in code.splitlines():
        paragraph = doc.add_paragraph(style="Mermaid Code")
        set_paragraph_shading(paragraph, "F4F4F4")
        if line:
            paragraph.add_run(line)
        else:
            paragraph.add_run(" ")


def build_doc():
    markdown = SOURCE.read_text(encoding="utf-8")
    sections = parse_mermaid_sections(markdown)
    if not sections:
        raise RuntimeError("No Mermaid diagram sections found.")

    doc = Document()
    setup_styles(doc)
    add_title_block(doc)

    intro = doc.add_paragraph()
    intro.add_run(
        "Each section below contains the Mermaid source for one diagram. "
        "Select the grey code block and copy it into a Mermaid-enabled editor, "
        "Markdown document, or diagram renderer."
    )

    for index, (title, code) in enumerate(sections, start=1):
        doc.add_heading(f"{index}. {title}", level=1)
        note = doc.add_paragraph(style="Small Note")
        note.add_run("Copyable Mermaid source:")
        add_copyable_code_block(doc, code)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
