from pathlib import Path
import math

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.shared import Inches, Pt


WORKSPACE = Path(r"C:\Users\Edv\Documents\New project")
OUT_DIR = WORKSPACE / "ANONYMUS_diagram_images"
DOCX_PATH = WORKSPACE / "ANONYMUS_diagram_images.docx"


COLORS = {
    "bg": "#F8FAFC",
    "ink": "#111827",
    "muted": "#4B5563",
    "line": "#475569",
    "blue": "#2563EB",
    "blue_light": "#DBEAFE",
    "blue_soft": "#EFF6FF",
    "green": "#16A34A",
    "green_light": "#DCFCE7",
    "green_soft": "#F0FDF4",
    "rose": "#DB2777",
    "rose_light": "#FCE7F3",
    "amber": "#D97706",
    "amber_light": "#FEF3C7",
    "purple": "#7C3AED",
    "purple_light": "#F3E8FF",
    "slate_light": "#E2E8F0",
    "white": "#FFFFFF",
}


def font(size, bold=False, mono=False):
    candidates = []
    if mono:
        candidates = [
            r"C:\Windows\Fonts\consola.ttf",
            r"C:\Windows\Fonts\cour.ttf",
        ]
    elif bold:
        candidates = [
            r"C:\Windows\Fonts\arialbd.ttf",
            r"C:\Windows\Fonts\calibrib.ttf",
        ]
    else:
        candidates = [
            r"C:\Windows\Fonts\arial.ttf",
            r"C:\Windows\Fonts\calibri.ttf",
        ]

    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


TITLE = font(46, bold=True)
SUBTITLE = font(25)
GROUP_TITLE = font(27, bold=True)
BOX_TITLE = font(25, bold=True)
BOX_TEXT = font(20)
SMALL = font(17)
CLASS_TITLE = font(23, bold=True)
CLASS_TEXT = font(17, mono=True)
OVAL_TEXT = font(19, bold=True)


def text_size(draw, text, fnt):
    bbox = draw.textbbox((0, 0), text, font=fnt)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def wrap_text(draw, text, fnt, max_width):
    wrapped = []
    for raw_line in text.split("\n"):
        words = raw_line.split()
        if not words:
            wrapped.append("")
            continue
        line = words[0]
        for word in words[1:]:
            candidate = f"{line} {word}"
            if text_size(draw, candidate, fnt)[0] <= max_width:
                line = candidate
            else:
                wrapped.append(line)
                line = word
        wrapped.append(line)
    return wrapped


def center_text(draw, box, text, fnt, fill=COLORS["ink"], line_gap=6):
    x, y, w, h = box
    lines = wrap_text(draw, text, fnt, w - 28)
    heights = [text_size(draw, line or " ", fnt)[1] for line in lines]
    total_h = sum(heights) + line_gap * (len(lines) - 1)
    yy = y + (h - total_h) / 2
    for line, height in zip(lines, heights):
        tw, _ = text_size(draw, line, fnt)
        draw.text((x + (w - tw) / 2, yy), line, font=fnt, fill=fill)
        yy += height + line_gap


def left_text(draw, x, y, text, fnt, fill=COLORS["ink"], max_width=None, line_gap=5):
    lines = text.split("\n") if max_width is None else wrap_text(draw, text, fnt, max_width)
    yy = y
    for line in lines:
        draw.text((x, yy), line, font=fnt, fill=fill)
        yy += text_size(draw, line or " ", fnt)[1] + line_gap
    return yy


def rounded_box(draw, xy, fill, outline, width=3, radius=24):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def label_box(draw, x, y, w, h, title, body="", fill=COLORS["white"], outline=COLORS["line"], accent=COLORS["blue"]):
    shadow = (x + 7, y + 7, x + w + 7, y + h + 7)
    draw.rounded_rectangle(shadow, radius=20, fill="#D1D5DB")
    rounded_box(draw, (x, y, x + w, y + h), fill, outline, width=3, radius=20)
    draw.rounded_rectangle((x, y, x + w, y + 48), radius=20, fill=accent)
    draw.rectangle((x, y + 24, x + w, y + 48), fill=accent)
    tw, th = text_size(draw, title, BOX_TITLE)
    draw.text((x + (w - tw) / 2, y + 13), title, font=BOX_TITLE, fill=COLORS["white"])
    if body:
        center_text(draw, (x + 10, y + 54, w - 20, h - 60), body, BOX_TEXT, fill=COLORS["ink"])


def group_box(draw, x, y, w, h, title, fill, outline):
    rounded_box(draw, (x, y, x + w, y + h), fill, outline, width=4, radius=28)
    draw.text((x + 28, y + 20), title, font=GROUP_TITLE, fill=outline)


def arrow(draw, start, end, color=COLORS["line"], width=4, label=None, label_offset=(0, 0), dashed=False):
    x1, y1 = start
    x2, y2 = end
    if dashed:
        dash_len = 20
        gap = 12
        dist = math.hypot(x2 - x1, y2 - y1)
        steps = max(1, int(dist / (dash_len + gap)))
        for i in range(steps + 1):
            a = i / (steps + 1)
            b = min(1, a + dash_len / dist)
            sx = x1 + (x2 - x1) * a
            sy = y1 + (y2 - y1) * a
            ex = x1 + (x2 - x1) * b
            ey = y1 + (y2 - y1) * b
            draw.line((sx, sy, ex, ey), fill=color, width=width)
    else:
        draw.line((x1, y1, x2, y2), fill=color, width=width)

    angle = math.atan2(y2 - y1, x2 - x1)
    size = 18
    p1 = (x2, y2)
    p2 = (x2 - size * math.cos(angle - math.pi / 6), y2 - size * math.sin(angle - math.pi / 6))
    p3 = (x2 - size * math.cos(angle + math.pi / 6), y2 - size * math.sin(angle + math.pi / 6))
    draw.polygon([p1, p2, p3], fill=color)

    if label:
        mx = (x1 + x2) / 2 + label_offset[0]
        my = (y1 + y2) / 2 + label_offset[1]
        tw, th = text_size(draw, label, SMALL)
        draw.rounded_rectangle((mx - tw / 2 - 8, my - th / 2 - 5, mx + tw / 2 + 8, my + th / 2 + 5), radius=8, fill=COLORS["white"])
        draw.text((mx - tw / 2, my - th / 2), label, font=SMALL, fill=color)


def poly_arrow(draw, points, color=COLORS["line"], width=4, label=None, label_at=0.5, dashed=False):
    if len(points) < 2:
        return
    for start, end in zip(points, points[1:]):
        x1, y1 = start
        x2, y2 = end
        if dashed:
            dist = math.hypot(x2 - x1, y2 - y1)
            if dist == 0:
                continue
            dash_len = 20
            gap = 12
            steps = max(1, int(dist / (dash_len + gap)))
            for i in range(steps + 1):
                a = i / (steps + 1)
                b = min(1, a + dash_len / dist)
                sx = x1 + (x2 - x1) * a
                sy = y1 + (y2 - y1) * a
                ex = x1 + (x2 - x1) * b
                ey = y1 + (y2 - y1) * b
                draw.line((sx, sy, ex, ey), fill=color, width=width)
        else:
            draw.line((x1, y1, x2, y2), fill=color, width=width)

    x1, y1 = points[-2]
    x2, y2 = points[-1]
    angle = math.atan2(y2 - y1, x2 - x1)
    size = 18
    p1 = (x2, y2)
    p2 = (x2 - size * math.cos(angle - math.pi / 6), y2 - size * math.sin(angle - math.pi / 6))
    p3 = (x2 - size * math.cos(angle + math.pi / 6), y2 - size * math.sin(angle + math.pi / 6))
    draw.polygon([p1, p2, p3], fill=color)

    if label:
        # Put the label on the middle segment by default.
        segment_index = max(0, min(len(points) - 2, int((len(points) - 1) * label_at)))
        sx, sy = points[segment_index]
        ex, ey = points[segment_index + 1]
        mx = (sx + ex) / 2
        my = (sy + ey) / 2
        tw, th = text_size(draw, label, SMALL)
        draw.rounded_rectangle((mx - tw / 2 - 8, my - th / 2 - 5, mx + tw / 2 + 8, my + th / 2 + 5), radius=8, fill=COLORS["white"])
        draw.text((mx - tw / 2, my - th / 2), label, font=SMALL, fill=color)


def title(draw, w, main, sub):
    tw, _ = text_size(draw, main, TITLE)
    draw.text(((w - tw) / 2, 34), main, font=TITLE, fill=COLORS["ink"])
    sw, _ = text_size(draw, sub, SUBTITLE)
    draw.text(((w - sw) / 2, 90), sub, font=SUBTITLE, fill=COLORS["muted"])


def build_architecture():
    w, h = 2600, 1650
    img = Image.new("RGB", (w, h), COLORS["bg"])
    draw = ImageDraw.Draw(img)
    title(draw, w, "ANONYMUS System Architecture Diagram", "Flutter client, local device resources, and backend integrations")

    group_box(draw, 470, 165, 1120, 1270, "Flutter Client Application", COLORS["blue_soft"], COLORS["blue"])
    group_box(draw, 1650, 245, 410, 640, "Local Device", "#FFF7ED", COLORS["amber"])
    group_box(draw, 1650, 945, 720, 455, "Backend / External Services", COLORS["green_soft"], COLORS["green"])

    # User actor.
    draw.ellipse((120, 540, 210, 630), outline=COLORS["ink"], width=5)
    draw.line((165, 630, 165, 790), fill=COLORS["ink"], width=5)
    draw.line((90, 680, 240, 680), fill=COLORS["ink"], width=5)
    draw.line((165, 790, 95, 900), fill=COLORS["ink"], width=5)
    draw.line((165, 790, 235, 900), fill=COLORS["ink"], width=5)
    center_text(draw, (55, 925, 260, 70), "User / Reporter", BOX_TITLE)

    label_box(draw, 590, 250, 380, 115, "Entry Point", "main.dart\nrunApp(WomenSafetyApp)", accent=COLORS["blue"])
    label_box(draw, 1060, 250, 410, 115, "App Shell", "MaterialApp\nsession bootstrap", accent=COLORS["blue"])
    label_box(draw, 590, 455, 380, 130, "Settings Scope", "Inherited settings\nlocalization provider", accent=COLORS["purple"])
    label_box(draw, 1060, 455, 410, 130, "Theme + Strings", "AppTheme, AppPalette\nAppStrings", accent=COLORS["purple"])
    label_box(draw, 570, 705, 410, 155, "Onboarding UI", "WelcomePage\nSignInPage\nRegistrationPage", accent=COLORS["rose"])
    label_box(draw, 1060, 705, 430, 155, "Reporting UI", "ReportHomePage\nhome, report, SOS,\noffline, sync, settings", accent=COLORS["rose"])
    label_box(draw, 570, 1015, 410, 170, "Auth Services", "AuthApiService\nGoogleAuthService\nAuthSessionStore", accent=COLORS["blue"])
    label_box(draw, 1060, 1015, 430, 170, "Reporting Services", "ReportingApiService\nOfflineReportStore\nReportingSeedData", accent=COLORS["blue"])
    label_box(draw, 820, 1245, 430, 150, "Domain Models", "AuthSession, User, Category,\nLocationType, PendingReport,\nSubmissionResult", accent=COLORS["line"])

    label_box(draw, 1710, 345, 300, 100, "SharedPreferences", "session, settings,\noffline queue", accent=COLORS["amber"])
    label_box(draw, 1710, 535, 300, 100, "Geolocator", "current latitude\nand longitude", accent=COLORS["amber"])
    label_box(draw, 1710, 725, 300, 100, "Clipboard", "copy SOS message", accent=COLORS["amber"])

    label_box(draw, 1725, 1030, 315, 100, "Google Sign-In", "external identity provider", accent=COLORS["green"])
    label_box(draw, 2070, 1030, 285, 100, "REST API", "ApiConfig.baseUrl\n/api/v1", accent=COLORS["green"])
    label_box(draw, 1725, 1210, 315, 110, "Auth Endpoints", "/auth/register/\n/auth/sign-in/\n/auth/google/", accent=COLORS["green"])
    label_box(draw, 2070, 1210, 285, 110, "Report Endpoints", "/health/\n/taxonomies/*\n/reports/", accent=COLORS["green"])

    arrow(draw, (250, 735), (470, 735), label="uses app")
    arrow(draw, (970, 308), (1060, 308))
    arrow(draw, (1265, 365), (1265, 455))
    arrow(draw, (970, 520), (1060, 520))
    arrow(draw, (780, 585), (780, 705))
    arrow(draw, (1265, 585), (1265, 705))
    arrow(draw, (775, 860), (775, 1015))
    arrow(draw, (1275, 860), (1275, 1015))
    arrow(draw, (980, 1090), (1060, 1090), label="models", label_offset=(0, -34))
    arrow(draw, (1035, 1185), (1035, 1255))

    poly_arrow(draw, [(980, 1070), (1580, 1070), (1580, 400), (1710, 400)], label="session")
    poly_arrow(draw, [(1490, 1120), (1620, 1120), (1620, 430), (1710, 430)], label="offline queue")
    poly_arrow(draw, [(1490, 740), (1620, 740), (1620, 585), (1710, 585)], label="location")
    poly_arrow(draw, [(1490, 815), (1620, 815), (1620, 775), (1710, 775)], label="SOS copy")
    poly_arrow(draw, [(980, 1045), (1650, 1045), (1725, 1080)], label="Google auth")
    poly_arrow(draw, [(980, 1160), (2050, 1160), (2070, 1100)], label="auth HTTP")
    poly_arrow(draw, [(1490, 1095), (2050, 1095), (2070, 1080)], label="report HTTP")
    arrow(draw, (2210, 1130), (1875, 1210), label="auth routes")
    arrow(draw, (2210, 1130), (2210, 1210), label="report routes")

    path = OUT_DIR / "ANONYMUS_system_architecture_diagram.png"
    img.save(path, quality=95)
    return path


def draw_class(draw, x, y, w, title_text, lines, accent=COLORS["blue"]):
    header_h = 42
    line_h = 24
    h = header_h + 16 + max(1, len(lines)) * line_h + 16
    rounded_box(draw, (x, y, x + w, y + h), COLORS["white"], accent, width=3, radius=16)
    draw.rounded_rectangle((x, y, x + w, y + header_h), radius=16, fill=accent)
    draw.rectangle((x, y + 20, x + w, y + header_h), fill=accent)
    tw, _ = text_size(draw, title_text, CLASS_TITLE)
    draw.text((x + (w - tw) / 2, y + 10), title_text, font=CLASS_TITLE, fill=COLORS["white"])
    yy = y + header_h + 10
    for line in lines:
        draw.text((x + 14, yy), line, font=CLASS_TEXT, fill=COLORS["ink"])
        yy += line_h
    return (x, y, w, h)


def center_of(box):
    x, y, w, h = box
    return x + w / 2, y + h / 2


def build_class_diagram():
    w, h = 3400, 2450
    img = Image.new("RGB", (w, h), COLORS["bg"])
    draw = ImageDraw.Draw(img)
    title(draw, w, "ANONYMUS Class Diagram", "Main Flutter classes, services, models, and relationships")

    group_box(draw, 70, 165, 835, 980, "App Core", COLORS["blue_soft"], COLORS["blue"])
    group_box(draw, 960, 165, 720, 980, "Onboarding + Auth", COLORS["rose_light"], COLORS["rose"])
    group_box(draw, 1735, 165, 860, 980, "Reporting", COLORS["green_soft"], COLORS["green"])
    group_box(draw, 2645, 165, 680, 980, "Settings + Config", COLORS["purple_light"], COLORS["purple"])
    group_box(draw, 70, 1215, 3255, 1060, "Domain Models", COLORS["white"], COLORS["line"])

    women = draw_class(draw, 125, 255, 360, "WomenSafetyApp", [
        "- settingsController",
        "- sessionStore",
        "- session?",
        "+ bootstrapApp()",
        "+ handleAuthenticated()",
        "+ handleLoggedOut()",
        "+ build()",
    ])
    scope = draw_class(draw, 535, 255, 330, "AppSettingsScope", [
        "+ controllerOf()",
        "+ readControllerOf()",
        "+ stringsOf()",
        "+ readStringsOf()",
    ], COLORS["purple"])
    welcome = draw_class(draw, 125, 640, 360, "WelcomePage", [
        "+ onAuthenticated",
        "+ build()",
        "+ open sign up",
        "+ open sign in",
    ], COLORS["rose"])
    report_page = draw_class(draw, 535, 640, 330, "ReportHomePage", [
        "+ currentUser",
        "+ onLogout",
        "+ create state",
    ], COLORS["green"])

    sign_in = draw_class(draw, 1015, 255, 300, "SignInPage", [
        "- AuthApiService",
        "- GoogleAuthService",
        "+ submit()",
        "+ continueWithGoogle()",
    ], COLORS["rose"])
    register = draw_class(draw, 1345, 255, 300, "RegistrationPage", [
        "- AuthApiService",
        "- GoogleAuthService",
        "+ submit()",
        "+ continueWithGoogle()",
    ], COLORS["rose"])
    auth_api = draw_class(draw, 1015, 620, 300, "AuthApiService", [
        "+ register()",
        "+ signIn()",
        "+ signInWithGoogle()",
        "+ isConnectivityError()",
    ])
    google_auth = draw_class(draw, 1345, 620, 300, "GoogleAuthService", [
        "- GoogleSignIn",
        "- isInitialized",
        "+ authenticate()",
    ])
    session_store = draw_class(draw, 1180, 900, 330, "AuthSessionStore", [
        "+ loadSession()",
        "+ saveSession()",
        "+ clearSession()",
    ])

    state = draw_class(draw, 1795, 255, 440, "ReportHomePageState", [
        "- ReportingApiService",
        "- OfflineReportStore",
        "- selectedCategory",
        "- selectedLocationType",
        "- pendingReports",
        "+ useCurrentLocation()",
        "+ activateSosSupport()",
        "+ submit()",
        "+ queueCurrentReportOffline()",
        "+ syncPendingReports()",
    ], COLORS["green"])
    reporting_api = draw_class(draw, 2285, 255, 330, "ReportingApiService", [
        "+ isBackendAvailable()",
        "+ fetchCategories()",
        "+ fetchLocationTypes()",
        "+ submitReport()",
        "+ isConnectivityError()",
    ])
    offline = draw_class(draw, 2285, 645, 330, "OfflineReportStore", [
        "+ loadPendingReports()",
        "+ enqueueReport()",
        "+ savePendingReports()",
    ])
    seed = draw_class(draw, 1795, 900, 440, "ReportingSeedData", [
        "+ incidentCategories",
        "+ locationTypes",
        "offline taxonomy fallback",
    ], COLORS["line"])

    settings = draw_class(draw, 2705, 255, 560, "AppSettingsController", [
        "- language",
        "- themeMode",
        "- themePreset",
        "- backendUrl",
        "- autoSyncEnabled",
        "+ load()",
        "+ setLanguage()",
        "+ setThemeMode()",
        "+ setBackendUrl()",
    ], COLORS["purple"])
    strings = draw_class(draw, 2705, 650, 260, "AppStrings", [
        "+ text()",
        "+ categoryName()",
        "+ locationTypeName()",
        "+ statusLabel()",
    ], COLORS["purple"])
    theme = draw_class(draw, 2995, 650, 270, "AppTheme", [
        "+ light()",
        "+ dark()",
        "uses AppPalette",
    ], COLORS["purple"])
    config = draw_class(draw, 2705, 930, 560, "ApiConfig / GoogleAuthConfig", [
        "+ baseUrl",
        "+ normalizeBaseUrl()",
        "+ serverClientId",
        "+ isConfigured",
    ], COLORS["purple"])

    # Domain model rows.
    auth_user = draw_class(draw, 135, 1310, 390, "AuthenticatedUser", [
        "+ id",
        "+ fullName",
        "+ email",
        "+ toJson()",
        "+ fromJson()",
    ], COLORS["line"])
    auth_session = draw_class(draw, 575, 1310, 360, "AuthSession", [
        "+ token",
        "+ user",
        "+ toJson()",
        "+ fromJson()",
        "+ fromAuthResponse()",
    ], COLORS["line"])
    google_result = draw_class(draw, 985, 1310, 350, "GoogleAuthResult", [
        "+ idToken",
        "+ email",
        "+ displayName?",
    ], COLORS["line"])
    cat = draw_class(draw, 1385, 1310, 390, "IncidentCategory", [
        "+ id",
        "+ code",
        "+ name",
        "+ description?",
        "+ sortOrder",
        "+ fromJson()",
    ], COLORS["line"])
    loc = draw_class(draw, 1825, 1310, 360, "LocationType", [
        "+ id",
        "+ code",
        "+ name",
        "+ description?",
        "+ sortOrder",
        "+ fromJson()",
    ], COLORS["line"])
    pending = draw_class(draw, 2225, 1310, 520, "PendingIncidentReport", [
        "+ localId",
        "+ categoryCode",
        "+ locationTypeCode",
        "+ occurredAt",
        "+ latitude / longitude",
        "+ description",
        "+ consentAcknowledged",
        "+ queuedAt",
        "+ toJson() / fromJson()",
    ], COLORS["line"])
    result = draw_class(draw, 2795, 1310, 450, "ReportSubmissionResult", [
        "+ id",
        "+ publicReference",
        "+ status",
        "+ message",
        "+ fromJson()",
        "+ offlineQueued()",
    ], COLORS["line"])

    ex = draw_class(draw, 135, 1880, 510, "Exceptions", [
        "AuthApiException",
        "ReportingApiException",
        "GoogleAuthException",
        "GoogleAuthNotConfiguredException",
        "GoogleAuthCancelledException",
        "GoogleAuthFailedException",
    ], COLORS["amber"])
    enums = draw_class(draw, 735, 1880, 470, "Enums", [
        "AppLanguage: english, swahili",
        "AppThemePreset: roseDawn,",
        "oceanCalm, emeraldGlow",
    ], COLORS["amber"])

    # Relationships kept intentionally sparse so the exported image stays readable.
    arrow(draw, (485, 315), (535, 315), label="uses")
    arrow(draw, (305, 450), (305, 640), label="shows")
    arrow(draw, (485, 720), (535, 720), label="opens")
    arrow(draw, (865, 315), (2705, 315), label="owns settings")
    arrow(draw, (485, 760), (1015, 315), label="opens")
    arrow(draw, (485, 790), (1345, 315), label="opens")

    arrow(draw, (1165, 435), (1165, 620), label="calls")
    arrow(draw, (1495, 435), (1495, 620), label="calls")
    arrow(draw, (1315, 705), (1345, 705), label="uses")
    arrow(draw, (1330, 790), (1330, 900), label="stores")
    poly_arrow(draw, [(1315, 650), (1360, 650), (1360, 1220), (755, 1310)], label="returns")
    arrow(draw, (755, 1440), (525, 1440), label="has user")
    poly_arrow(draw, [(1510, 980), (1510, 1220), (755, 1220), (755, 1310)], label="loads")
    arrow(draw, (1495, 780), (1160, 1310), label="returns")

    arrow(draw, (865, 720), (1795, 340), label="creates")
    arrow(draw, (2235, 410), (2285, 410), label="uses")
    arrow(draw, (2235, 565), (2285, 720), label="queues")
    arrow(draw, (2015, 860), (2015, 900), label="fallback")
    poly_arrow(draw, [(2445, 835), (2445, 1200), (2485, 1310)], label="stores")
    poly_arrow(draw, [(2285, 360), (2090, 360), (2090, 1220), (1580, 1310)], label="loads")
    poly_arrow(draw, [(2380, 360), (2380, 1220), (2005, 1310)], label="loads")
    poly_arrow(draw, [(2615, 410), (3090, 410), (3090, 1310)], label="returns")
    arrow(draw, (1795, 965), (1580, 1510), label="seed")
    arrow(draw, (2050, 965), (2005, 1510), label="seed")

    arrow(draw, (2985, 565), (2985, 650), label="strings")
    arrow(draw, (2965, 720), (2995, 720), label="theme")
    arrow(draw, (2985, 930), (2985, 870), label="config")
    poly_arrow(draw, [(2705, 1045), (2600, 1045), (2600, 410), (2615, 410)], label="base URL", dashed=True)

    legend_x, legend_y = 2470, 2025
    draw.rounded_rectangle((legend_x, legend_y, legend_x + 810, legend_y + 170), radius=18, fill="#F8FAFC", outline=COLORS["line"], width=2)
    draw.text((legend_x + 20, legend_y + 18), "Relationship Legend", font=BOX_TITLE, fill=COLORS["ink"])
    draw.text((legend_x + 20, legend_y + 60), "solid arrow = direct dependency/use", font=BOX_TEXT, fill=COLORS["muted"])
    draw.text((legend_x + 20, legend_y + 95), "dashed arrow = configuration dependency", font=BOX_TEXT, fill=COLORS["muted"])
    draw.text((legend_x + 20, legend_y + 130), "lower panel = data/domain classes used by services and pages", font=BOX_TEXT, fill=COLORS["muted"])

    path = OUT_DIR / "ANONYMUS_class_diagram.png"
    img.save(path, quality=95)
    return path


def actor(draw, x, y, label):
    draw.ellipse((x + 55, y, x + 135, y + 80), outline=COLORS["ink"], width=5)
    draw.line((x + 95, y + 80, x + 95, y + 220), fill=COLORS["ink"], width=5)
    draw.line((x + 25, y + 130, x + 165, y + 130), fill=COLORS["ink"], width=5)
    draw.line((x + 95, y + 220, x + 30, y + 330), fill=COLORS["ink"], width=5)
    draw.line((x + 95, y + 220, x + 160, y + 330), fill=COLORS["ink"], width=5)
    center_text(draw, (x, y + 350, 190, 70), label, BOX_TITLE)


def oval(draw, x, y, w, h, text, fill=COLORS["white"], outline=COLORS["blue"]):
    draw.ellipse((x, y, x + w, y + h), fill=fill, outline=outline, width=3)
    center_text(draw, (x + 12, y + 8, w - 24, h - 16), text, OVAL_TEXT)
    return (x, y, w, h)


def build_usecase():
    w, h = 2700, 1650
    img = Image.new("RGB", (w, h), COLORS["bg"])
    draw = ImageDraw.Draw(img)
    title(draw, w, "ANONYMUS Use Case Diagram", "User actions and external systems for the Women Safety app")

    actor(draw, 95, 575, "User /\nReporter")

    group_box(draw, 390, 170, 1580, 1295, "Women Safety App", COLORS["blue_soft"], COLORS["blue"])

    auth_cases = [
        oval(draw, 500, 280, 330, 95, "View welcome\nscreen", COLORS["white"], COLORS["rose"]),
        oval(draw, 890, 280, 330, 95, "Register", COLORS["white"], COLORS["rose"]),
        oval(draw, 1280, 280, 330, 95, "Sign in", COLORS["white"], COLORS["rose"]),
        oval(draw, 500, 430, 330, 95, "Continue with\nGoogle", COLORS["white"], COLORS["rose"]),
        oval(draw, 890, 430, 330, 95, "Restore saved\nsession", COLORS["white"], COLORS["rose"]),
        oval(draw, 1280, 430, 330, 95, "Log out", COLORS["white"], COLORS["rose"]),
    ]

    report_cases = [
        oval(draw, 500, 665, 330, 95, "View home\ndashboard", COLORS["white"], COLORS["green"]),
        oval(draw, 890, 665, 330, 95, "Load categories\nand locations", COLORS["white"], COLORS["green"]),
        oval(draw, 1280, 665, 330, 95, "Create incident\nreport", COLORS["white"], COLORS["green"]),
        oval(draw, 500, 815, 330, 95, "Attach current\nlocation", COLORS["white"], COLORS["green"]),
        oval(draw, 890, 815, 330, 95, "Submit report\nonline", COLORS["white"], COLORS["green"]),
        oval(draw, 1280, 815, 330, 95, "Save report\noffline", COLORS["white"], COLORS["green"]),
        oval(draw, 500, 965, 330, 95, "View offline\nqueue", COLORS["white"], COLORS["green"]),
        oval(draw, 890, 965, 330, 95, "Sync saved\nreports", COLORS["white"], COLORS["green"]),
        oval(draw, 1280, 965, 330, 95, "Copy SOS\nmessage", COLORS["white"], COLORS["green"]),
    ]

    settings_cases = [
        oval(draw, 500, 1210, 330, 95, "View safety\nguide", COLORS["white"], COLORS["purple"]),
        oval(draw, 890, 1210, 330, 95, "Change app\nsettings", COLORS["white"], COLORS["purple"]),
        oval(draw, 1280, 1210, 330, 95, "Select theme\npreset", COLORS["white"], COLORS["purple"]),
        oval(draw, 1630, 280, 300, 95, "Configure\nbackend URL", COLORS["white"], COLORS["purple"]),
    ]

    # External systems.
    label_box(draw, 2115, 285, 360, 110, "Google Sign-In", "identity provider", accent=COLORS["green"])
    label_box(draw, 2115, 550, 360, 140, "REST Backend API", "auth, taxonomy,\nhealth, reports", accent=COLORS["green"])
    label_box(draw, 2115, 840, 360, 130, "Device Services", "location permission,\nGPS, clipboard", accent=COLORS["amber"])
    label_box(draw, 2115, 1135, 360, 130, "Local Storage", "session, settings,\noffline reports", accent=COLORS["amber"])

    # User lines to the main interaction groups.
    arrow(draw, (285, 795), (500, 327), width=3, label="starts")
    arrow(draw, (285, 795), (500, 715), width=3, label="reports")
    arrow(draw, (285, 795), (500, 1015), width=3, label="reviews")
    arrow(draw, (285, 795), (500, 1255), width=3, label="learns")
    arrow(draw, (285, 795), (890, 1255), width=3, label="configures")

    # Includes and dependencies inside the app.
    arrow(draw, (830, 327), (890, 327), label="may")
    arrow(draw, (1220, 327), (1280, 327), label="or")
    arrow(draw, (830, 477), (890, 477), label="stores session", dashed=True)
    arrow(draw, (830, 715), (890, 715), label="needs")
    arrow(draw, (1220, 715), (1280, 715), label="enables")
    arrow(draw, (1430, 760), (1060, 815), label="include", dashed=True)
    arrow(draw, (1430, 760), (1430, 815), label="include", dashed=True)
    arrow(draw, (1220, 865), (1280, 865), label="offline fallback", dashed=True)
    arrow(draw, (830, 1015), (890, 1015), label="then")
    arrow(draw, (1220, 1255), (1280, 1255), label="theme")

    # External system connections, grouped and routed to avoid crossing.
    poly_arrow(draw, [(830, 475), (1995, 475), (2115, 340)], label="OAuth")
    poly_arrow(draw, [(1055, 327), (1985, 327), (1985, 610), (2115, 610)], label="auth HTTP")
    poly_arrow(draw, [(1445, 715), (1985, 715), (1985, 610), (2115, 610)], label="taxonomy/report HTTP")
    poly_arrow(draw, [(1055, 1015), (1985, 1015), (1985, 610), (2115, 610)], label="sync HTTP")
    poly_arrow(draw, [(665, 865), (1975, 865), (2115, 905)], label="GPS")
    poly_arrow(draw, [(1610, 1015), (1975, 1015), (1975, 905), (2115, 905)], label="copy")
    poly_arrow(draw, [(1055, 477), (1995, 477), (1995, 1195), (2115, 1195)], label="session")
    poly_arrow(draw, [(1430, 865), (1995, 865), (1995, 1195), (2115, 1195)], label="offline queue")
    poly_arrow(draw, [(1055, 1255), (1995, 1255), (1995, 1195), (2115, 1195)], label="settings")
    poly_arrow(draw, [(1930, 327), (1995, 327), (1995, 1195), (2115, 1195)], label="save URL")

    path = OUT_DIR / "ANONYMUS_use_case_diagram.png"
    img.save(path, quality=95)
    return path


def build_docx(image_paths):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    title_p = doc.add_paragraph()
    title_run = title_p.add_run("ANONYMUS Project Diagram Images")
    title_run.bold = True
    title_run.font.size = Pt(24)
    doc.add_paragraph("The following pages contain the diagrams as image objects. You can copy the images directly from this document or use the PNG files in the image folder.")

    for idx, path in enumerate(image_paths, start=1):
        doc.add_page_break()
        p = doc.add_paragraph()
        run = p.add_run(path.stem.replace("_", " ").title())
        run.bold = True
        run.font.size = Pt(16)
        doc.add_picture(str(path), width=Inches(10.0))

    doc.save(DOCX_PATH)


def main():
    OUT_DIR.mkdir(exist_ok=True)
    image_paths = [
        build_architecture(),
        build_class_diagram(),
        build_usecase(),
    ]
    build_docx(image_paths)
    for path in image_paths:
        print(path)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
